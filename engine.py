# -*- coding: utf-8 -*-
"""
KJFM 논문 빌더 — 로컬 웹앱
칸에 붙여넣기 → '워드로 내보내기' → 검증된 KJFM 템플릿 스타일로 .docx 생성.

원리: 실제 게재본 템플릿(1.안동균_final.docx)을 열어 본문만 비우고
      페이지 설정·테마·명명 스타일은 그대로 유지한 뒤, 입력 내용을
      해당 스타일(2. 본문 / 1. 대제목_국문 ...)로 새로 채운다. → 양식이 깨질 수 없음.
"""
import os
import re
import io
import copy
import json
import base64
import urllib.parse

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.text.paragraph import Paragraph
from docx.table import Table

HERE = os.path.dirname(os.path.abspath(__file__))
# 템플릿은 스킬 안에 번들(assets/template.docx). 없으면 원본 OneDrive 경로로 폴백.
_BUNDLED = os.path.join(HERE, 'assets', 'template.docx')
_FALLBACK = os.environ.get("KJFM_FALLBACK_TEMPLATE", "")
TEMPLATE = _BUNDLED if os.path.isfile(_BUNDLED) else _FALLBACK

# 입력 칸 → 템플릿 스타일 매핑
S_TITLE_EN = "Title_English"
S_TITLE_KO = "국문제목"
S_AUTHOR   = "차례제"
S_AUTHOR_NAME = "저자명"      # 저자명 줄(HY신명조 9pt)
S_DATE_BLANK  = "9.게재일"    # 게재일 앞 빈 줄 스타일
S_ABSTRACT = "국문 본문"
S_H1       = "1. 대제목_국문"
S_H2       = "1.1. 중제목_국문"
S_H3       = "1.1.1. 소제목_국문"
S_H4       = "1.1.1.1. 세부제목_국문"
S_BODY     = "2. 본문"
S_BLANK    = "국문 본문"    # 절제목 앞 빈 줄
S_REF_HEAD = "References head"
S_REF      = "10. Refer"
S_TBL_CAP  = "3. 표"        # 표 제목(위, 좌측 9pt)
S_TBL_CELL = "7. 표 내용"   # 표 셀 내용(8pt 가운데)
S_FIG_CAP  = "4. 그림"      # 그림 제목(아래, 가운데 9pt 굵게)
S_ABS_TITLE = "abstract title"            # 초록 박스: "Abstract" 제목
S_ABS_TEXT  = "Abstract_text"             # 초록 박스: 본문/키워드(국)/JEL
S_ABS_KW    = "Keyword(TimesNewRoman9pt)" # 초록 박스: 영문 Keywords

HEADING_STYLES = {'h1': S_H1, 'h2': S_H2, 'h3': S_H3, 'h4': S_H4}

COPYRIGHT_TEXT = ("ⓒ Copyright: The Author(s). This is an Open Access article distributed under the "
                  "terms of the Creative Commons Attribution Non-Commercial License "
                  "(https://creativecommons.org/licenses/by-nc/4.0/) which permits unrestricted "
                  "non-commercial use, distribution, and reproduction in any medium, provided the "
                  "original work is properly cited.")

_WNS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def _add_footnote_ref(paragraph, fid, mark=None, half_pt='16', superscript=True):
    """단락에 각주 참조 추가 — TNR. 기본 8pt 위첨자. mark 지정 시 번호 없는 커스텀 마크.
    half_pt: 글자 크기(반포인트, 8pt=16 / 15.5pt=31). superscript: 위첨자 여부.
    마크 텍스트는 footnoteReference와 같은 run 안에 둔다(Word 호환)."""
    r = paragraph.add_run()
    rpr = r._r.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman'); rpr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), half_pt); rpr.append(sz)
    if superscript:
        va = OxmlElement('w:vertAlign'); va.set(qn('w:val'), 'superscript'); rpr.append(va)
    ref = OxmlElement('w:footnoteReference'); ref.set(qn('w:id'), str(fid))
    if mark:
        ref.set(qn('w:customMarkFollows'), '1')
    r._r.append(ref)
    if mark:                                  # 마크 텍스트(ⓒ)를 같은 run에, ref 바로 뒤에
        t = OxmlElement('w:t'); t.text = mark; r._r.append(t)


def _inject_footnotes(docx_bytes, items):
    """word/footnotes.xml의 내용 각주(id>=1)를 items로 교체. items: [(id, text), ...]."""
    if not items:
        return docx_bytes
    import zipfile
    from lxml import etree
    zin = zipfile.ZipFile(io.BytesIO(docx_bytes))
    names = zin.namelist()
    if 'word/footnotes.xml' not in names:
        return docx_bytes
    root = etree.fromstring(zin.read('word/footnotes.xml'))
    for fn in root.findall(_WNS + 'footnote'):
        fid = fn.get(_WNS + 'id')
        if fid is not None and int(fid) >= 1:
            root.remove(fn)
    for item in items:
        fid, text = item[0], item[1]
        mark = item[2] if len(item) > 2 else None
        fn = etree.SubElement(root, _WNS + 'footnote'); fn.set(_WNS + 'id', str(fid))
        p = etree.SubElement(fn, _WNS + 'p')
        ppr = etree.SubElement(p, _WNS + 'pPr')
        ind = etree.SubElement(ppr, _WNS + 'ind'); ind.set(_WNS + 'left', '255'); ind.set(_WNS + 'hanging', '255')
        # 마커: 숫자 각주는 footnoteRef, 커스텀 마크는 마크 텍스트(ⓒ)
        r1 = etree.SubElement(p, _WNS + 'r')
        rpr1 = etree.SubElement(r1, _WNS + 'rPr')
        rf1 = etree.SubElement(rpr1, _WNS + 'rFonts'); rf1.set(_WNS + 'ascii', 'Times New Roman'); rf1.set(_WNS + 'hAnsi', 'Times New Roman')
        etree.SubElement(rpr1, _WNS + 'sz').set(_WNS + 'val', '16')
        if mark:
            mt = etree.SubElement(r1, _WNS + 't'); mt.text = mark
        else:
            etree.SubElement(rpr1, _WNS + 'vertAlign').set(_WNS + 'val', 'superscript')
            etree.SubElement(r1, _WNS + 'footnoteRef')
        # 본문 텍스트: Times New Roman 8pt
        r2 = etree.SubElement(p, _WNS + 'r')
        rpr2 = etree.SubElement(r2, _WNS + 'rPr')
        rf2 = etree.SubElement(rpr2, _WNS + 'rFonts')
        rf2.set(_WNS + 'ascii', 'Times New Roman'); rf2.set(_WNS + 'hAnsi', 'Times New Roman'); rf2.set(_WNS + 'eastAsia', '맑은 고딕')
        etree.SubElement(rpr2, _WNS + 'sz').set(_WNS + 'val', '16')
        t = etree.SubElement(r2, _WNS + 't')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        # 커스텀 마크면 마크(ⓒ)가 r1에 있으니 본문에서 선두 ⓒ 제거. 각주도 * * *→***, 0.X→.X
        text = _norm_sig_zero(text)
        t.text = ' ' + (text.lstrip('ⓒ© ').strip() if mark else text)
    new_fn = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
    out = io.BytesIO()
    zout = zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED)
    for n in names:
        zout.writestr(n, new_fn if n == 'word/footnotes.xml' else zin.read(n))
    zout.close()
    return out.getvalue()


# ---------------------------------------------------------------- 기호 간격 교정 (선택)
def normalize_spacing(text):
    """KJFM 기호 간격 규칙(안전판). 본문/초록 텍스트에만 적용.
    하이픈 단어(well-being)·DOI·URL은 건드리지 않는다."""
    if not text:
        return text
    t = text
    # 괄호: '(' 앞 공백, ')' 뒤 공백
    t = re.sub(r'(?<=\S)\(', ' (', t)
    t = re.sub(r'\)(?=[^\s\)\.,;:?!])', ') ', t)
    # 이항 기호 앞뒤 공백
    for sym in ['<', '>', '=', '~', '+', '×', '±', '≥', '≤', '→']:
        t = t.replace(sym, ' ' + sym + ' ')
    # 숫자 범위 하이픈(3-4 → 3 - 4), 단어 하이픈(well-being)은 유지
    t = re.sub(r'(?<=\d)\s*[-–~]\s*(?=\d)', ' - ', t)
    # 다중 공백 → 1개 (줄바꿈 보존)
    t = re.sub(r'[ \t]{2,}', ' ', t)
    # 구두점 앞 공백 제거
    t = re.sub(r'\s+([,.\;:?!])', r'\1', t)
    return t.strip()


_SIG_SPACE_RE = re.compile(r'(?<=\*)[ \t]+(?=\*)')
_LEADING_ZERO_RE = re.compile(r'(?<![\w.])(-?)0\.(\d)')


def _norm_sig_zero(text):
    """모든 출력 공통: 유의수준 별표 사이 공백 제거('* * *'→'***') + 선행 0 제거(0.XXX→.XXX).
    DOI/연도(10.x, 2020.)는 앞이 숫자라 영향 없음."""
    if not text:
        return text
    text = _SIG_SPACE_RE.sub('', text)
    text = _LEADING_ZERO_RE.sub(r'\1.\2', text)
    return text


def normalize_doi(text):
    """franchise APA 7th: 다양한 DOI 표기 → https://doi.org/10... 통일."""
    if not text:
        return text
    t = re.sub(r'https?\s*://\s*(dx\.)?doi\.org\s*/\s*', 'https://doi.org/', text, flags=re.I)
    t = re.sub(r'\bdoi\s*[:：]\s*(?=10\.)', 'https://doi.org/', t, flags=re.I)
    return t


# ---------------------------------------------------------------- docx 생성
def open_template_skeleton():
    """템플릿을 열어 ①상단 ISSN/DOI 표 ②초록 박스(1x1) 골격을 복사로 확보하고,
    본문을 비운다(sectPr·스타일·테마·머리말 유지). 반환: (doc, sectPr, top표el, 초록박스el)."""
    doc = Document(TEMPLATE)
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    tables = doc.tables
    top_el = copy.deepcopy(tables[0]._tbl) if len(tables) >= 1 else None
    abs_el = copy.deepcopy(tables[1]._tbl) if len(tables) >= 2 else None
    for child in list(body):
        if child is not sectPr:
            body.remove(child)
    return doc, sectPr, top_el, abs_el


def _add_copyright_footer(section):
    """1페이지 하단 푸터에 저작권 문구 — 각주 아님, Arial 8pt 내어쓰기(소속 각주와 같은 글씨체)."""
    section.different_first_page_header_footer = True
    footer = section.first_page_footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for r in list(para.runs):
        r.text = ''
    pf = para.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.left_indent = Cm(0.45)
    pf.first_line_indent = Cm(-0.45)   # 내어쓰기(© Copyright가 매달림)
    pf.line_spacing = 1.0
    run = para.add_run(COPYRIGHT_TEXT)
    run.font.size = Pt(8)
    run.font.name = 'Arial'


# ─────────────────────────── 머릿글(running head) ───────────────────────────
def _running_authors(authors):
    """저자 영문명만 추출(한글·역할 제거): 'Heejun LEE 이희준(주), …' → 'Heejun LEE, Dongkyun AHN'."""
    names = []
    for part in re.split(r'[,，]', authors or ''):
        m = re.match(r"\s*([A-Za-z][A-Za-z.'\- ]*[A-Za-z])", part)
        if m:
            names.append(re.sub(r'\s+', ' ', m.group(1)).strip())
    return ', '.join(names)


def _doi_parts(doi):
    """DOI '…2026.6.17.2.1'에서 (연도, 권, 호) 추정 — 연.월.권.호.순번."""
    nums = re.findall(r'\d+', doi or '')
    if len(nums) >= 5:
        y, _mo, vol, iss, _seq = nums[-5:]
        return y, vol, iss
    return '', '', ''


def _running_odd_text(data):
    """홀수 페이지 머릿글: '저자 / Korean Journal of Franchise Management 권-호 (연도) 페이지'."""
    au = _running_authors(data.get('authors', ''))
    vi = (data.get('vol_issue', '') or '').strip()
    yr = (data.get('year', '') or '').strip()
    pg = (data.get('pages', '') or '').strip()
    if not (vi and yr):
        y, vol, iss = _doi_parts(data.get('doi', ''))
        if not vi and vol and iss:
            vi = '%s(%s)' % (vol, iss)        # 권(호) 형식: 17(2)
        if not yr and y:
            yr = y
    tail = 'Korean Journal of Franchise Management'
    if vi:
        tail += ' ' + vi
    if yr:
        tail += ' (%s)' % yr
    if pg:
        tail += ' ' + pg
    return ('%s / %s' % (au, tail)) if au else tail


def _running_even_text(data):
    """짝수 페이지 머릿글: 영문 논문제목(여러 줄 → 한 줄)."""
    return re.sub(r'\s*\n\s*', ' ', (data.get('title_en', '') or '').strip()).strip()


def _add_page_field(para):
    """현재 페이지 번호 필드(PAGE) run 추가."""
    run = para.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = ' PAGE '
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    return run


def _style_run_tnr(run):
    """Times New Roman 8pt 이탤릭."""
    run.font.name = 'Times New Roman'; run.font.size = Pt(8); run.font.italic = True
    rpr = run._r.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.insert(0, rf)
    rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman')
    rf.set(qn('w:eastAsia'), 'Times New Roman'); rf.set(qn('w:cs'), 'Times New Roman')


def _hdr_borderless(tbl):
    """머릿글 표 테두리 제거."""
    tblPr = tbl._tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for e in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + e); el.set(qn('w:val'), 'nil'); b.append(el)
    tblPr.append(b)


def _hdr_cell(cell, width, align):
    cell.width = Emu(int(width))
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.paragraph_format.alignment = align
    p.paragraph_format.space_after = Pt(0); p.paragraph_format.space_before = Pt(0)
    return p


def _fill_running_header(header, kind, text, tw):
    """표(테두리 없음) 머릿글 — 긴 텍스트는 칸 안에서 자동 줄바꿈, 쪽번호는 안 밀림.
    odd: [본문 가운데][쪽번호 오른쪽] / even: [쪽번호 왼쪽][본문 가운데]."""
    header.is_linked_to_previous = False
    for p in header.paragraphs:                           # 기존 단락 비움
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        p.paragraph_format.space_after = Pt(0)
    tbl = header.add_table(rows=1, cols=2, width=Emu(int(tw)))
    tbl.autofit = False
    _hdr_borderless(tbl)
    page_w = Cm(1.3); text_w = int(tw) - int(page_w)
    c0, c1 = tbl.rows[0].cells
    if kind == 'odd':
        _style_run_tnr(_hdr_cell(c0, text_w, WD_ALIGN_PARAGRAPH.CENTER).add_run(text))
        _style_run_tnr(_add_page_field(_hdr_cell(c1, page_w, WD_ALIGN_PARAGRAPH.RIGHT)))
    else:
        _style_run_tnr(_add_page_field(_hdr_cell(c0, page_w, WD_ALIGN_PARAGRAPH.LEFT)))
        _style_run_tnr(_hdr_cell(c1, text_w, WD_ALIGN_PARAGRAPH.CENTER).add_run(text))
    hdr_el = tbl._tbl.getparent()                         # 표를 맨 앞으로(빈 단락은 뒤로)
    hdr_el.remove(tbl._tbl); hdr_el.insert(0, tbl._tbl)


def _set_page_num_start(section, n):
    sectPr = section._sectPr
    pgt = sectPr.find(qn('w:pgNumType'))
    if pgt is None:
        pgt = OxmlElement('w:pgNumType'); sectPr.append(pgt)
    pgt.set(qn('w:start'), str(n))


def _add_running_heads(doc, data):
    """홀/짝 다른 머릿글 + 홀수 시작 페이지. 모든 섹션에 적용."""
    if not data.get('running_head', True):
        return
    settings = doc.settings.element
    if settings.find(qn('w:evenAndOddHeaders')) is None:
        settings.append(OxmlElement('w:evenAndOddHeaders'))
    odd = _running_odd_text(data)
    even = _running_even_text(data)
    try:
        start_n = int(re.sub(r'\D', '', str(data.get('start_page', '1'))) or '1')
    except Exception:
        start_n = 1
    if start_n < 1:
        start_n = 1
    if start_n % 2 == 0:        # 논문은 홀수 페이지로 시작
        start_n += 1
    _set_page_num_start(doc.sections[0], start_n)
    for sec in doc.sections:
        tw = sec.page_width - sec.left_margin - sec.right_margin
        _fill_running_header(sec.header, 'odd', odd, tw)
        _fill_running_header(sec.even_page_header, 'even', even, tw)


def _set_section_cols(section, num, space=425):
    """섹션의 단(컬럼) 수 설정."""
    sp = section._sectPr
    cols = sp.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        sp.append(cols)
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), str(space))


def _fill_abstract_box(doc, abs_el, abstract, kw_en, kw_ko, jel, fix=False):
    """복사한 초록 박스(1x1 표)의 셀 내용을 새로 채운다(테두리·스타일 유지)."""
    tbl = Table(abs_el, doc)
    cell = tbl.cell(0, 0)
    for p in list(cell.paragraphs):           # 기존 셀 단락 제거
        p._p.getparent().remove(p._p)

    avail = [s.name for s in doc.styles]

    def addp(text, style, bold_label=False, bold_all=False):
        p = cell.add_paragraph()
        if style in avail:
            p.style = doc.styles[style]
        # 'Purpose:', 'Result:' 등 라벨(콜론 앞)만 굵게
        m = re.match(r'^([A-Za-z가-힣][A-Za-z가-힣 ,/&\-]{1,45}?[:：])(.*)$', text) if bold_label else None
        if m:
            r1 = p.add_run(m.group(1)); r1.bold = True
            p.add_run(m.group(2))
        else:
            r = p.add_run(text)
            if bold_all:
                r.bold = True
        return p

    addp('Abstract', S_ABS_TITLE)
    # 초록: 주제별 줄바꿈 없이 한 단락으로 이어지게. 구조 라벨(Purpose:/Result: 등)만 인라인 굵게.
    abs_lines = [l.strip() for l in (abstract or '').replace('\r\n', '\n').split('\n') if l.strip()]
    if abs_lines:
        full = _norm_sig_zero(' '.join((normalize_spacing(l) if fix else l) for l in abs_lines))
        ap = cell.add_paragraph()
        if S_ABS_TEXT in avail:
            ap.style = doc.styles[S_ABS_TEXT]
        labels = re.compile(
            r'(Purpose|Research design[^:：]{0,45}|Methodology|Results?|Findings?|'
            r'Conclusions?|Implications?|Originality[^:：]{0,45}|연구목적|연구방법|연구결과|결론)\s*[:：]', re.I)
        pos = 0
        for m in labels.finditer(full):
            if m.start() > pos:
                ap.add_run(full[pos:m.start()])
            ap.add_run(m.group(0)).bold = True
            pos = m.end()
        if pos < len(full):
            ap.add_run(full[pos:])
    # 키워드: 라벨(Keywords:/키워드:)만 굵게, 나머지 보통 — Keyword(TimesNewRoman9pt) 스타일
    if (kw_en or '').strip():
        addp('Keywords: ' + kw_en.strip(), S_ABS_KW, bold_label=True)
    if (kw_ko or '').strip():
        addp('키워드: ' + kw_ko.strip(), S_ABS_KW, bold_label=True)
    if (jel or '').strip():
        addp('JEL Classification Code: ' + jel.strip(), S_ABS_TEXT)
    addp('', S_ABS_TEXT)   # JEL 아래 빈 줄 1개


_HYP_RE = re.compile(
    r'^\s*(H\s?\d+[a-z]?(?:[-.]\d+)?|Hypothesis\s+\d+[a-z]?|가설\s*\d+[a-z]?)\s*([:：.)])\s*(.*)$', re.I)


def add_lines(doc, text, style, fix=False, force_hyp=False):
    """text를 줄 단위로 나눠 각 줄을 style 단락으로 추가. 빈 줄은 건너뜀.
    모든 줄에 _norm_sig_zero('* * *'→'***', 0.X→.X) 적용.
    가설: H1:/H2:/Hypothesis/가설 줄(또는 force_hyp=가설블록이면 모든 줄)은 내어쓰기,
    라벨('H1:') 굵게, 가설 묶음 앞뒤에 빈 줄."""
    if not text:
        return
    in_hyp = [False]

    def close_hyp():
        if in_hyp[0]:
            doc.add_paragraph('', style=S_BLANK)   # 가설 묶음 뒤 빈 줄
            in_hyp[0] = False

    for line in text.replace('\r\n', '\n').split('\n'):
        line = line.strip()
        if not line:
            continue
        if fix:
            line = normalize_spacing(line)
        line = _norm_sig_zero(line)                # 모든 줄: * * *→***, 0.X→.X
        m = _HYP_RE.match(line)
        if m or force_hyp:
            if not in_hyp[0]:
                doc.add_paragraph('', style=S_BLANK)   # 가설 묶음 앞 빈 줄
                in_hyp[0] = True
            p = doc.add_paragraph(style=style)
            pf = p.paragraph_format
            pf.left_indent = Cm(0.7); pf.first_line_indent = Cm(-0.7)   # 내어쓰기
            if m:
                lab = p.add_run(m.group(1) + m.group(2)); lab.bold = True   # 'H1:' 굵게
                if m.group(3):
                    p.add_run(' ' + m.group(3))
            else:
                p.add_run(line)                    # 라벨 없는 줄도 내어쓰기(가설 블록)
        else:
            close_hyp()
            doc.add_paragraph(line, style=style)
    close_hyp()


def add_figure(doc, block):
    """그림: 이미지(가운데, 폭 제한) + 캡션(아래, '4. 그림')."""
    img = block.get('image', '')
    if img:
        if img.strip().startswith('data:') and ',' in img:
            img = img.split(',', 1)[1]
        try:
            raw = base64.b64decode(img)
            shape = doc.add_picture(io.BytesIO(raw))
            maxw = Cm(15)
            if shape.width > maxw:  # 너무 넓으면 비율 유지하며 축소
                ratio = maxw / shape.width
                shape.height = int(shape.height * ratio)
                shape.width = maxw
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            doc.add_paragraph('[그림을 불러올 수 없습니다]', style=S_BODY)
    cap = block.get('caption', '').strip()
    if cap:
        doc.add_paragraph('', style=S_BLANK)            # 그림 아래 한 칸(엔터)
        cp = doc.add_paragraph(cap, style=S_TBL_CAP)    # 표 제목 형식, 가운데 정렬
        cp.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_caption_note(doc, block.get('note', ''))


def _add_caption_note(doc, note):
    """표/그림 아래 주석(Note) 줄 — 8pt, 'Note.' 라벨만 이탤릭."""
    note = _norm_sig_zero((note or '').strip())       # * * *→***, 0.X→.X
    if not note:
        return
    p = doc.add_paragraph(style=S_REF)
    pf = p.paragraph_format
    pf.left_indent = Cm(0); pf.first_line_indent = Cm(0)
    pf.space_before = Pt(0); pf.space_after = Pt(0)   # 표 바로 밑(엔터 없이)
    m = re.match(r'^(Note|주|주의)\s*[.:：]?\s*', note)
    if m:
        lab = p.add_run(note[:m.end()]); lab.font.size = Pt(8); lab.font.italic = True
        rest = p.add_run(note[m.end():]); rest.font.size = Pt(8)
    else:
        p.add_run(note).font.size = Pt(8)


HDR_FILL = 'B3B3B3'   # 표 제목 행: 흑백 70%(밝기) 음영
SUB_FILL = 'D1D1D1'   # 소제목(그룹머리) 행: 흑백 82% 음영


def _shade_cell(cell, hex_fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)


def _norm_table_cell(s):
    """표 셀 정규화: 공통 _norm_sig_zero('* * *'→'***', 0.XXX→.XXX)."""
    return _norm_sig_zero((s or '').strip())


def _cell_set_bottom(cell):
    """셀 아래 가로선(헤더 행 밑줄)."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(old)
    tb = OxmlElement('w:tcBorders')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '0'); bottom.set(qn('w:color'), '000000')
    tb.append(bottom); tcPr.append(tb)


def _apa_table_borders(table):
    """학술 3선 표: 표 위·아래 + 헤더 행 아래 가로선만, 세로선·내부선 없음(PDF·인쇄 적합)."""
    tblPr = table._tbl.tblPr
    for old in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(old)
    b = OxmlElement('w:tblBorders')
    for e in ('top', 'bottom'):
        el = OxmlElement('w:' + e)
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '8')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), '000000'); b.append(el)
    for e in ('left', 'right', 'insideH', 'insideV'):
        el = OxmlElement('w:' + e); el.set(qn('w:val'), 'nil'); b.append(el)
    tblPr.append(b)
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            _cell_set_bottom(cell)


def add_table(doc, block):
    """표: 캡션(위, '3. 표') + 탭구분 데이터로 Word 표 생성.
    3선 학술표(격자 없음)·셀 텍스트 가운데·헤더/소제목 볼드+음영. 0.XXX→.XXX, '* * *'→'***'."""
    cap = block.get('caption', '').strip()
    if cap:
        doc.add_paragraph(cap, style=S_TBL_CAP)
    data = block.get('data', '')
    rows = [r.split('\t') for r in data.replace('\r\n', '\n').split('\n') if r.strip() != '']
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncol)
    cell_style = doc.styles[S_TBL_CELL] if S_TBL_CELL in [s.name for s in doc.styles] else None
    for i, r in enumerate(rows):
        nonempty = [j for j in range(ncol) if j < len(r) and r[j].strip()]
        is_header = (i == 0)
        is_sub = (not is_header) and nonempty == [0]   # 첫 칸만 채워진 행 = 소제목(그룹머리)
        for j in range(ncol):
            cell = table.cell(i, j)
            cell.text = _norm_table_cell(r[j]) if j < len(r) else ''
            for p in cell.paragraphs:
                if cell_style is not None:
                    p.style = cell_style
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER   # 셀 텍스트 가운데
                if is_header or is_sub:
                    for run in p.runs:
                        run.bold = True
            if is_header:
                _shade_cell(cell, HDR_FILL)
            elif is_sub:
                _shade_cell(cell, SUB_FILL)
    _apa_table_borders(table)                         # 3선 표(격자 제거)
    _add_caption_note(doc, block.get('note', ''))   # 표 아래 Note란
    doc.add_paragraph('', style=S_BODY)  # 표 뒤 간격


def _apply_style_indent(doc, p, style_name):
    """명명 스타일의 들여쓰기를 단락에 직접 복사 — docx-preview/PDF가 스타일 들여쓰기를
    무시해도 내어쓰기(hanging indent)가 렌더되게."""
    try:
        st = doc.styles[style_name].paragraph_format
    except KeyError:
        return
    if st.left_indent is not None:
        p.paragraph_format.left_indent = st.left_indent
    if st.first_line_indent is not None:
        p.paragraph_format.first_line_indent = st.first_line_indent


def _add_reference(doc, line):
    """참고문헌 1줄 → 10. Refer 단락. franchise APA: DOI 통일 + 학술지명·권 이탤릭.
    권(volume) 위치를 찾아 그 앞 '. ' 이후(학술지명)부터 권 숫자까지만 이탤릭."""
    line = _norm_sig_zero(normalize_doi(line))       # * * *→***, 0.X→.X (DOI/연도 영향 없음)
    p = doc.add_paragraph(style=S_REF)
    _apply_style_indent(doc, p, S_REF)               # 내어쓰기 직접 적용(PDF 렌더용)
    s = e = None
    m = re.search(r',\s*(\d+)\s*[\(,]', line)   # ", 38(" 또는 ", 38,"
    if m:
        vol_end = m.start(1) + len(m.group(1))
        dot = line[:m.start()].rfind('. ')
        if dot != -1:
            s, e = dot + 2, vol_end
    if s is not None and s < e:
        p.add_run(line[:s])
        p.add_run(line[s:e]).italic = True
        p.add_run(line[e:])
    else:
        p.add_run(line)
    return p


def build_review_line(received, revised, accepted):
    parts = []
    if received.strip():
        parts.append(f"Received: {received.strip()}")
    if revised.strip():
        parts.append(f"Revised: {revised.strip()}")
    if accepted.strip():
        parts.append(f"Accepted: {accepted.strip()}")
    return "    ".join(parts)


def build_docx(data):
    """학회지 골격(상단 ISSN/DOI 표 + 초록 박스)을 유지하며 내용을 채워 .docx 생성.
    순서: 상단표 → 영문제목 → 국문제목 → 저자 → 심사일 → 초록박스 → 본문 → 참고문헌."""
    fix = bool(data.get('fix_spacing'))
    doc, sectPr, top_el, abs_el = open_template_skeleton()

    def add_el(el):  # sectPr 바로 앞에 요소 삽입(= 현재 본문 끝에 추가)
        if el is not None:
            sectPr.addprevious(el)

    # 1) 상단 ISSN/DOI 표(학회지 헤더 블록) — DOI 치환
    doi = (data.get('doi', '') or '').strip()
    if top_el is not None and doi:
        full = doi if doi.lower().startswith('http') else ('http://dx.doi.org/10.21871/KJFM.' + doi.lstrip('.'))
        tbl = Table(top_el, doc)
        for row in tbl.rows:
            for cell in row.cells:
                if 'doi.org' in cell.text.lower():
                    for p in cell.paragraphs:
                        if 'doi.org' in p.text.lower():
                            for r in p.runs:
                                r.text = ''
                            (p.runs[0] if p.runs else p.add_run()).text = full
    add_el(top_el)
    # ISSN 표 아래 8pt 빈 줄(제목 위)
    doc.add_paragraph().add_run().font.size = Pt(8)
    # 2) 영문 제목 / 3) 국문 제목
    add_lines(doc, data.get('title_en', ''), S_TITLE_EN)
    title_last_p = doc.paragraphs[-1] if (data.get('title_en', '') or '').strip() else None
    add_lines(doc, data.get('title_ko', ''), S_TITLE_KO)
    # 4) 저자 (9.5pt) — 저자별 소속 각주(¹²) 부착
    fn_items = []
    fid = [0]

    def next_fid():
        fid[0] += 1
        return fid[0]

    # 제목 설명각주 — ⚠️ '*' 커스텀마크는 Word가 제목을 깨뜨림 → 안전한 번호 각주 사용.
    # 단, 제목의 마커 크기만 15.5pt TNR로(요청), 본문 각주 텍스트는 8pt.
    if (data.get('title_note', '') or '').strip() and title_last_p is not None:
        _tnid = next_fid()
        _add_footnote_ref(title_last_p, _tnid, half_pt='31')   # 번호 각주, 마커 15.5pt 위첨자
        fn_items.append((_tnid, data['title_note'].strip()))

    affs = [a.strip() for a in (data.get('affiliations', '') or '').split('\n') if a.strip()]
    authors_line = (data.get('authors', '') or '').strip()
    if authors_line:
        p = doc.add_paragraph(style=S_AUTHOR_NAME)   # 저자명: HY신명조 9pt(앞 14pt)
        parts = [a.strip() for a in authors_line.split(',') if a.strip()]
        for i, ap in enumerate(parts):
            p.add_run(ap)
            if i < len(affs):
                _id = next_fid()
                _add_footnote_ref(p, _id)
                fn_items.append((_id, affs[i]))
            if i < len(parts) - 1:
                p.add_run(', ')
    # 5) 심사일 — 앞에 '게재일' 스타일 빈 줄 + 심사일(8pt) + 라이선스 각주(2단 각주영역 우측)
    review = build_review_line(data.get('received', ''), data.get('revised', ''), data.get('accepted', ''))
    anchor = None
    if review:
        doc.add_paragraph('', style=S_DATE_BLANK)    # 게재일 앞 빈 줄
        anchor = doc.add_paragraph(style=S_AUTHOR)
        anchor.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER   # 게재일 가운데 정렬
        anchor.paragraph_format.left_indent = Cm(0); anchor.paragraph_format.first_line_indent = Cm(0)
        anchor.add_run(review).font.size = Pt(8)
    if data.get('copyright', True):
        if anchor is None:
            anchor = doc.add_paragraph(style=S_AUTHOR)
        _cid = next_fid()
        _add_footnote_ref(anchor, _cid)              # 라이선스 각주(안전한 번호 — ⓒ마크는 Word 거부)
        fn_items.append((_cid, COPYRIGHT_TEXT))
    # 6) 초록 박스(영문 초록 + 영/국 키워드 + JEL) — 박스 골격 유지
    if abs_el is not None and (data.get('abstract_en', '').strip() or data.get('keywords_en', '').strip()
                               or data.get('keywords_ko', '').strip() or data.get('jel', '').strip()):
        _fill_abstract_box(doc, abs_el, data.get('abstract_en', ''), data.get('keywords_en', ''),
                           data.get('keywords_ko', ''), data.get('jel', ''), fix=fix)
        add_el(abs_el)

    # ── 단(컬럼) 전환: 헤더(상단표·제목·저자·초록박스)=1단 → 본문=2단 ──
    two_col = data.get('two_column', True)
    _set_section_cols(doc.sections[-1], 1)
    body_sec = None
    if two_col:
        body_sec = doc.add_section(WD_SECTION.CONTINUOUS)
        _set_section_cols(body_sec, 2)

    # 7) 본문 블록. 절제목 앞 빈 줄: 대제목 2·중제목 1 / 뒤 빈 줄: 모두 1.
    #    인접한 제목 사이 빈 줄은 '최댓값'으로 합쳐 누적(겹침) 방지.
    BEFORE = {'h1': 2, 'h2': 1, 'h3': 1, 'h4': 1}
    AFTER = {'h1': 1, 'h2': 1, 'h3': 1, 'h4': 1}
    pending = [0]      # 다음 콘텐츠 앞에 넣을 빈 줄 수(누적 max)
    seen_h1 = [False]

    def flush_blanks():
        for _ in range(pending[0]):
            doc.add_paragraph('', style=S_BLANK)
        pending[0] = 0

    for block in data.get('blocks', []):
        btype = block.get('type', 'body')
        if btype in HEADING_STYLES:
            before = BEFORE[btype]
            if btype == 'h1' and not seen_h1[0]:
                before = 0                       # 첫 대제목은 섹션 시작이라 앞 빈줄 없음
            seen_h1[0] = seen_h1[0] or (btype == 'h1')
            pending[0] = max(pending[0], before)
            flush_blanks()
            add_lines(doc, block.get('text', ''), HEADING_STYLES[btype])
            pending[0] = AFTER[btype]            # 제목 뒤 빈 줄(다음 콘텐츠 전 flush)
        elif btype == 'figure':
            flush_blanks(); add_figure(doc, block)
        elif btype == 'table':
            flush_blanks(); add_table(doc, block)
        elif btype == 'hyp':                         # 가설 블록: 모든 줄을 가설 포맷(내어쓰기·라벨 굵게)
            flush_blanks(); add_lines(doc, block.get('text', ''), S_BODY, fix=fix, force_hyp=True)
        else:
            flush_blanks(); add_lines(doc, block.get('text', ''), S_BODY, fix=fix)

    # 8) 참고문헌 — 대제목(숫자 없이) 'References' + franchise APA(DOI 통일·학술지명/권 이탤릭)
    refs = data.get('references', '').strip()
    if refs:
        pending[0] = max(pending[0], 2)
        flush_blanks()
        doc.add_paragraph('References', style=S_H1)   # 대제목, 번호 없음
        doc.add_paragraph('', style=S_BLANK)
        for ln in refs.split('\n'):
            if ln.strip():
                _add_reference(doc, ln.strip())

    # 9) 부록(Appendix) — 새 페이지(Ctrl+Enter) + 1단 레이아웃, 대제목 'Appendixes'(숫자 없이)
    appendix = data.get('appendix', [])
    if appendix:
        ap_sec = doc.add_section(WD_SECTION.NEW_PAGE)   # 다음 페이지로 넘김
        _set_section_cols(ap_sec, 1)                    # 1단 레이아웃
        doc.add_paragraph('Appendixes', style=S_H1)     # 대제목, 번호 없음
        doc.add_paragraph('', style=S_BLANK)
        _note_re = r'^(Note|주\b|주\s*[:：]|\*|α|β)'
        i = 0
        while i < len(appendix):
            item = appendix[i]; t = item.get('type')
            if t == 'table':
                note = ''                                # 표 바로 뒤 Note는 표의 주석으로 붙임(엔터 없이)
                if i + 1 < len(appendix) and appendix[i + 1].get('type') not in ('table', 'figure'):
                    nt = appendix[i + 1].get('text', '').strip()
                    if re.match(_note_re, nt):
                        note = nt; i += 1
                add_table(doc, {'caption': '', 'data': item.get('data', ''), 'note': note})
            elif t == 'figure':                          # 이미지로 된 부록 표/그림도 렌더(사라지던 버그)
                note = ''
                if i + 1 < len(appendix) and appendix[i + 1].get('type') not in ('table', 'figure'):
                    nt = appendix[i + 1].get('text', '').strip()
                    if re.match(_note_re, nt):
                        note = nt; i += 1
                add_figure(doc, {'caption': '', 'image': item.get('image', ''), 'note': note})
            else:
                txt = item.get('text', '').strip()
                if re.match(r'^(Appendix(es)?|부록)\s*[:：]?\s*$', txt, re.I):
                    i += 1; continue                    # 파일의 'Appendix' 단독 제목 제거(위 Appendixes와 중복)
                if re.match(r'^(Appendix(es|ix)?\s*\d|부록\s*\d)', txt, re.I):
                    doc.add_paragraph(txt, style=S_TBL_CAP)   # "Appendix 1: …" → 표 제목 스타일
                elif re.match(_note_re, txt):
                    p = doc.add_paragraph(txt, style=S_REF)   # 단독 Note — 작게
                    for r in p.runs:
                        r.font.size = Pt(8)
                else:
                    add_lines(doc, txt, S_REF)
            i += 1

    # 10) 머릿글(running head): 홀수=저자/학회지, 짝수=논문제목, TNR 8pt 이탤릭 + 홀수 시작 페이지
    _add_running_heads(doc, data)

    buf = io.BytesIO()
    doc.save(buf)
    return _inject_footnotes(buf.getvalue(), fn_items)


_MONTHS = {'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
           'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12}


def _yymmdd(s):
    """'May 10, 2026' → '260510'."""
    m = re.search(r'([A-Za-z]{3,9})\.?\s+(\d{1,2}),?\s+(\d{4})', s or '')
    if not m:
        return ''
    mon = _MONTHS.get(m.group(1)[:3].lower())
    if not mon:
        return ''
    return f'{int(m.group(3)) % 100:02d}{mon:02d}{int(m.group(2)):02d}'


def _filename_authors(authors):
    """저자 문자열에서 한글 이름(+역할 (주)/(공)/(교)) 추출 → '이도경(주) 민자경(교)'."""
    names = re.findall(r'[가-힣]{2,4}\s*(?:\([가-힣]\))?', authors or '')
    return ' '.join(n.replace(' ', '') for n in names)


def _vol_issue(data):
    """권(호) 문자열 '17(2)' — vol_issue 입력 우선, 없으면 DOI에서."""
    vi = (data.get('vol_issue', '') or '').strip()
    if vi:
        return vi
    y, vol, iss = _doi_parts(data.get('doi', ''))
    return ('%s(%s)' % (vol, iss)) if (vol and iss) else ''


_SOFFICE = None


def _find_soffice():
    """LibreOffice(soffice) 실행 경로 탐색(1회 캐시)."""
    global _SOFFICE
    if _SOFFICE is not None:
        return _SOFFICE
    import shutil
    for c in ['/Applications/LibreOffice.app/Contents/MacOS/soffice',
              shutil.which('soffice'), shutil.which('libreoffice'),
              '/opt/homebrew/bin/soffice', '/usr/bin/libreoffice']:
        if c and os.path.exists(c):
            _SOFFICE = c
            return c
    _SOFFICE = ''
    return ''


def _docx_to_pdf(docx_bytes):
    """docx 바이트 → PDF 바이트(LibreOffice headless). 실패 시 None."""
    soffice = _find_soffice()
    if not soffice:
        return None
    import tempfile, subprocess, glob as _glob, shutil
    d = tempfile.mkdtemp(prefix='kjfmpdf_')
    try:
        dx = os.path.join(d, 'doc.docx')
        with open(dx, 'wb') as f:
            f.write(docx_bytes)
        prof = 'file://' + os.path.join(d, 'profile')   # 호출별 프로필 → 동시 변환 가능
        subprocess.run([soffice, '-env:UserInstallation=' + prof, '--headless',
                        '--convert-to', 'pdf', '--outdir', d, dx],
                       capture_output=True, timeout=90)
        pdfs = _glob.glob(os.path.join(d, '*.pdf'))
        if pdfs:
            with open(pdfs[0], 'rb') as f:
                return f.read()
        return None
    except Exception:
        return None
    finally:
        shutil.rmtree(d, ignore_errors=True)


def _cite_keys_from_body(text):
    """본문에서 인용 (첫저자 성, 연도) 추출 — 괄호형은 ';'로 나눠 인용단위별 첫 저자만."""
    keys = set()
    for m in re.finditer(r'\(([^)]*\d{4}[^)]*)\)', text):              # (A, 2020; B & C, 2019)
        for unit in m.group(1).split(';'):
            if not re.search(r'(?:19|20)\d{2}', unit):
                continue
            sm = re.search(r'([A-Z][A-Za-z\-]+|[가-힣]{2,4})', unit)
            ym = re.search(r'((?:19|20)\d{2})', unit)
            if sm and ym:
                keys.add((sm.group(1).lower(), ym.group(1)))
    for m in re.finditer(r'([A-Z][A-Za-z\-]+|[가-힣]{2,4})'             # Author (2020) 서술형
                         r'(?:\s+(?:et al\.?|and\s+[A-Z][A-Za-z\-]+|&\s*[A-Z][A-Za-z\-]+|등|와|과))?'
                         r'\s*\(((?:19|20)\d{2})', text):
        keys.add((m.group(1).lower(), m.group(2)))
    return keys


def _ref_key(ref):
    """참고문헌 한 줄 → (첫 저자 성, 연도)."""
    ym = re.search(r'\(?((?:19|20)\d{2})[a-z]?\)?', ref)
    year = ym.group(1) if ym else ''
    sm = re.match(r'\s*([A-Z][A-Za-z\-]+|[가-힣]{2,4})', ref)
    sur = sm.group(1).lower() if sm else ''
    return (sur, year)


def check_citations(data):
    """본문 인용 ↔ 참고문헌 대조. 누락(인용했으나 목록 없음)·미인용(목록만 있고 본문 인용 없음) 보고."""
    parts = [data.get('abstract_en', '') or '']
    for b in data.get('blocks', []):
        t = b.get('type', '')
        if t == 'body' or t.startswith('h'):
            parts.append(b.get('text', '') or '')
        elif t in ('table', 'figure'):
            parts.append((b.get('caption', '') or '') + ' ' + (b.get('note', '') or ''))
    body = '\n'.join(parts)
    refs = [r.strip() for r in (data.get('references', '') or '').split('\n') if r.strip()]
    cites = _cite_keys_from_body(body)
    ref_keys = {}
    for r in refs:
        ref_keys.setdefault(_ref_key(r), r)

    def match(a, b):   # 성 접두 일치 + 연도 동일
        return a[1] == b[1] and a[1] and (a[0].startswith(b[0]) or b[0].startswith(a[0])) and a[0] and b[0]

    missing = sorted({'%s (%s)' % (s, y) for (s, y) in cites
                      if not any(match((s, y), rk) for rk in ref_keys)})
    uncited = [r[:70] for k, r in ref_keys.items()
               if not any(match(c, k) for c in cites)]
    return {'n_cites': len(cites), 'n_refs': len(refs),
            'missing': missing, 'uncited': sorted(uncited)}


def make_filename(data):
    """{순번}_KJFM_{권(호)}_{투고일YYMMDD}_{게재확정일YYMMDD}_{저자(역할) …}.docx"""
    seq = (data.get('seq', '') or '').strip()
    vi = _vol_issue(data)
    recv = _yymmdd(data.get('received', ''))
    acc = _yymmdd(data.get('accepted', ''))
    authors = _filename_authors(data.get('authors', ''))
    parts = [p for p in [seq, 'KJFM', vi, recv, acc, authors] if p]
    name = '_'.join(parts) if len(parts) > 1 else (
        (data.get('title_ko', '') or data.get('title_en', '') or 'KJFM_논문').split('\n')[0].strip())
    name = re.sub(r'[\\/:*?"<>|]', '', name)[:80] or 'KJFM_논문'
    return name + '.docx'


# ---------------------------------------------------------------- docx 자동 파싱 (불러오기)
def _has_hangul(s):
    return bool(re.search(r'[가-힣]', s))


_DATE_RE = re.compile(r'[A-Za-z]{3,9}\.?\s+\d{1,2},?\s+\d{4}')


def _parse_review_line(text):
    """'Received: Feb 12, 2026  Revised: ... Accepted: ...' → 날짜 3개로 분해."""
    out = {'received': '', 'revised': '', 'accepted': ''}
    labels = [('Received', 'received'), ('Revised', 'revised'), ('Accepted', 'accepted')]
    spans = []
    for label, key in labels:
        m = re.search(label + r'\s*[:：]?\s*', text)
        if m:
            spans.append((m.start(), m.end(), key))
    spans.sort()
    for i, (s, e, key) in enumerate(spans):
        nxt = spans[i + 1][0] if i + 1 < len(spans) else len(text)
        seg = text[e:nxt]
        dm = _DATE_RE.search(seg)
        out[key] = dm.group(0).strip() if dm else seg.strip(' \t,;　.')
    return out


def _rows_to_tsv(rows):
    """표 셀 2차원 리스트 → 탭/줄바꿈 구분 문자열(빌더 표 블록 형식)."""
    out = []
    for r in rows:
        out.append('\t'.join((c or '').replace('\t', ' ').replace('\n', ' ').strip() for c in r))
    return '\n'.join(out)


_A_BLIP = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
_R_EMBED = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'


def _iter_body_items(doc):
    """본문 자식을 문서 순서대로 순회 → ('p', style, text) / ('tbl', Table) / ('img', base64)."""
    body = doc.element.body
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            blips = child.findall('.//' + _A_BLIP)   # 단락 안 이미지(그림)
            if blips:
                rid = blips[0].get(_R_EMBED)
                try:
                    blob = doc.part.related_parts[rid].blob
                    import base64
                    yield ('img', base64.b64encode(blob).decode())
                    continue
                except Exception:
                    pass
            p = Paragraph(child, doc)
            yield ('p', (p.style.name if p.style else 'Normal'), p.text.strip())
        elif child.tag == qn('w:tbl'):
            yield ('tbl', Table(child, doc))


def _table_tsv(tbl):
    rows = [[c.text.strip() for c in r.cells] for r in tbl.rows]
    return _rows_to_tsv(rows)


def _table_cell_text(tbl):
    """표 전체 셀의 단락 텍스트를 줄바꿈 유지하여 결합(1x1 박스 추출용)."""
    parts = []
    for row in tbl.rows:
        for cell in row.cells:
            parts.append('\n'.join(p.text for p in cell.paragraphs))
    return '\n'.join(parts)


def _looks_like_abstract_box(tbl):
    """초록 박스(1x1, 'Abstract'/'Purpose:' 포함)인지."""
    if len(tbl.rows) != 1 or len(tbl.columns) != 1:
        return False
    t = _table_cell_text(tbl)
    return bool(re.search(r'\bAbstract\b', t) or re.search(r'Purpose\s*[:：]', t, re.I)) and len(t) > 120


def _extract_abstract_box(tbl, res):
    """초록 박스 셀 → abstract_en / keywords_en / keywords_ko / jel 로 분해."""
    _extract_abstract_text(_table_cell_text(tbl), res)


def _extract_abstract_text(text, res):
    """초록 텍스트 → abstract_en / keywords_en / keywords_ko / jel."""
    abs_lines = []
    for line in text.split('\n'):
        s = line.strip()
        if not s:
            continue
        if re.match(r'^Abstract$', s, re.I):
            continue
        mk = re.match(r'^(Keywords?|키워드|주제어)\s*[:：]\s*(.*)', s, re.I)
        if mk:
            val = mk.group(2).strip().rstrip('.').strip()
            if mk.group(1).lower().startswith('key'):
                res['keywords_en'] = val
            else:
                res['keywords_ko'] = val
            continue
        mj = re.match(r'^JEL.*?[:：]\s*(.*)', s, re.I)
        if mj:
            res['jel'] = mj.group(1).strip()
            continue
        abs_lines.append(s)
    if abs_lines and not res.get('abstract_en'):
        res['abstract_en'] = '\n'.join(abs_lines).strip()


_TITLE_STYLES = ('Title_English', '국문제목', 'Title', 'title')


def _is_old_header_table(tbl):
    """구형 포맷: 표 셀 안에 제목/저자가 스타일째 들어있는 헤더 표인지.
    초록 박스(Abstract/Purpose 1x1)는 제외 — 박스 안 키워드가 저자 스타일이어도 헤더표 아님."""
    if _looks_like_abstract_box(tbl):
        return False
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                sty = (p.style.name if p.style else '')
                if sty in ('Title_English', 'author', '저자명', '차례제', '2 저자성명') and p.text.strip():
                    return True
    return False


def _extract_header_table(tbl, res):
    """구형 헤더 표 셀에서 제목(en/ko)·저자·심사일·DOI·초록 추출."""
    title_en, title_ko, authors = [], [], []
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                t = p.text.strip()
                if not t:
                    continue
                sty = (p.style.name if p.style else '')
                low = sty.lower()
                if 'issn' in low or 'doi' in low:
                    m = re.search(r'KJFM\.([0-9.]+)', t)
                    if m and not res.get('doi'):
                        res['doi'] = m.group(1).rstrip('.')
                    continue
                if re.search(r'Received|Revised|Accepted', t):
                    for k, v in _parse_review_line(t).items():
                        if v:
                            res[k] = v
                    continue
                if sty in _AUTHOR_STYLES or _looks_like_author_line(t):
                    authors.append(t)
                    continue
                if 'abstract' in low or 'keyword' in low or sty in ('Abstract_text', 'abstract title'):
                    continue                      # 초록은 셀 전체로 따로 처리
                (title_ko if _has_hangul(t) else title_en).append(t)
    if title_en and not res.get('title_en'):
        res['title_en'] = '\n'.join(title_en)
    if title_ko and not res.get('title_ko'):
        res['title_ko'] = '\n'.join(title_ko)
    if authors and not res.get('authors'):
        res['authors'] = ' '.join(authors).strip()
    for row in tbl.rows:                          # 초록 셀
        for cell in row.cells:
            ct = '\n'.join(p.text for p in cell.paragraphs)
            if re.search(r'\bAbstract\b', ct) and len(ct) > 120:
                _extract_abstract_text(ct, res)
                return


def _table_looks_like_data(tsv):
    """데이터 표 판별: 2열 이상, 셀이 과하게 길지 않음(레이아웃·초록 박스 제외)."""
    rows = [r.split('\t') for r in tsv.split('\n') if r != '']
    if len(rows) < 2:
        return False
    ncol = max(len(r) for r in rows)
    maxlen = max((len(c) for r in rows for c in r), default=0)
    return ncol >= 2 and maxlen <= 400


_HEADING_ALIASES = ('1. 대제목_국문', '1.1. 중제목_국문', '중제목', '1.1.1. 소제목_국문',
                    '0.0. sub heading', '한1.1.1한', '1.1.1.1. 세부제목_국문', '0.0.0.sub heading')
_AUTHOR_STYLES = ('차례제', 'author', '저자명', '2 저자성명', 'Author')


def _looks_like_author_line(txt):
    """내용 기준 저자 줄 판별: 역할표기 (주)(교)(공) 또는 '영문이름 한글이름' 패턴(스타일 무관)."""
    if not txt or len(txt) > 140:
        return False
    if re.search(r'\(\s*[주교공]\s*\)', txt):
        return True
    # "Mi-Na KWAK 곽미나" 류 (영문 1~3토큰 + 한글 2~4자) 1쌍 이상
    pairs = re.findall(r'[A-Z][A-Za-z.\'\-]+(?:\s+[A-Z][A-Za-z.\'\-]+){0,2}\s+[가-힣]{2,4}', txt)
    return len(pairs) >= 1


def _parse_kjfm(items, res):
    """KJFM 명명 스타일 문서 → 위치+스타일 기준 분해 (표·부록 포함).
    제목/저자/심사일은 '초록 박스/첫 절제목 전'(헤더 영역)에서 스타일과 무관하게 위치로 판별."""
    title_en, title_ko, refs, abstract = [], [], [], []
    appendix = []
    pending = []
    pending_cap = ['', '']  # [표캡션, 그림캡션] 대기
    started = [False]       # 본문 시작 여부(헤더 영역 레이아웃 표 제외용)
    header_done = [False]   # 제목/저자/심사일 영역 끝남(초록 박스/첫 헤딩 후)
    in_refs = [False]       # 'References' 이후 = 참고문헌 영역(스타일 무관)
    in_appendix = [False]   # 'Appendix...' 이후 = 부록 영역
    res['appendix'] = appendix

    def flush():
        if pending:
            res['blocks'].append({'type': 'body', 'text': '\n'.join(pending)})
            pending.clear()

    def drop_orphan_caps():
        # 표/그림이 따라오지 않은 캡션은 본문으로 흡수(텍스트 보존)
        for c in pending_cap:
            if c:
                pending.append(c)
        pending_cap[0] = pending_cap[1] = ''

    for it in items:
        if it[0] == 'tbl':
            tbl = it[1]
            if in_appendix[0]:                          # 부록 표 → 순서대로 보존
                appendix.append({'type': 'table', 'data': _table_tsv(tbl)})
            elif not header_done[0] and _is_old_header_table(tbl):  # 구형: 표 안에 제목·저자·초록
                _extract_header_table(tbl, res)
                header_done[0] = True
            elif _looks_like_abstract_box(tbl):         # 초록 박스 → 초록/키워드/JEL 추출
                _extract_abstract_box(tbl, res)
                header_done[0] = True                   # 초록 박스 = 헤더 영역 끝
            elif started[0] and _table_looks_like_data(_table_tsv(tbl)):
                flush()
                res['blocks'].append({'type': 'table', 'caption': pending_cap[0], 'data': _table_tsv(tbl)})
            pending_cap[0] = ''
            continue
        if it[0] == 'img':                              # 이미지(그림) → 위치대로 그림 블록
            flush(); drop_orphan_caps()
            target = appendix if in_appendix[0] else res['blocks']
            target.append({'type': 'figure', 'caption': pending_cap[1], 'image': 'data:image/png;base64,' + it[1]})
            pending_cap[1] = ''
            continue
        sty, txt = it[1], it[2]
        if not txt:
            continue
        started[0] = True
        # ── 헤더 영역(제목/저자/심사일): 초록 박스/첫 절제목 전, 스타일 무관 위치 판별 ──
        if not header_done[0]:
            is_heading = sty in _HEADING_ALIASES or bool(re.match(r'^\d+\.\s+\S', txt))
            if is_heading or sty == 'References head':
                header_done[0] = True            # 본문 시작 → 아래 일반 처리로 흐름
            else:
                drop_orphan_caps()
                if re.search(r'Received|Revised|Accepted', txt):
                    for k, v in _parse_review_line(txt).items():
                        if v:
                            res[k] = v
                    header_done[0] = True        # 심사일 다음은 초록/본문 — 제목영역 종료(제목 누락 방지)
                elif re.match(r'^\s*(Keywords?|주제어|키워드|Abstract|초록|Purpose|Research\s+design|Findings|Conclusions?|JEL)\b', txt, re.I):
                    pass                         # 초록/키워드 줄은 제목·저자로 잡지 않음(스타일이 저자여도) — 박스/본문에서 처리
                elif sty in _AUTHOR_STYLES or _looks_like_author_line(txt):
                    res['authors'] = (res['authors'] + ' ' + txt).strip() if res['authors'] else txt
                else:                            # 제목(스타일 무관, 한글이면 국문칸)
                    (title_ko if _has_hangul(txt) else title_en).append(txt)
                continue
        # ── 참고문헌/부록 영역 (스타일 무관, 텍스트 기준 — 논문마다 스타일 다름) ──
        # 부록 시작은 References 없이도 감지(어디서든 'Appendix'/'부록'; 짧은 제목/캡션만)
        if (not in_appendix[0] and len(txt) < 80
                and (re.match(r'^(Appendix(es|ix)?|부록)\b', txt, re.I)
                     or re.match(r'^\d+\.?\s*Appendix(es|ix)?\b', txt, re.I))):
            drop_orphan_caps(); flush(); in_appendix[0] = True; in_refs[0] = False
            appendix.append({'type': 'cap', 'text': txt})
            continue
        if not in_refs[0] and not in_appendix[0] and (
                sty == 'References head' or re.match(r'^(References|참고\s*문헌)\s*$', txt, re.I)):
            drop_orphan_caps(); flush(); in_refs[0] = True
            continue
        if in_refs[0] or in_appendix[0]:
            if re.match(r'^(Appendix(es|ix)?|부록)\b', txt, re.I) or re.match(r'^\d+\.?\s*Appendix', txt, re.I):
                in_appendix[0] = True; in_refs[0] = False
                appendix.append({'type': 'cap', 'text': txt})
            elif in_appendix[0]:
                appendix.append({'type': 'cap', 'text': txt})
            elif not re.match(r'^(References|참고\s*문헌)\s*$', txt, re.I):
                refs.append(txt)                 # 참고문헌 항목(스타일 무관)
            continue
        # ── 본문 영역 ──
        if sty == '국문 본문':
            # 초록은 박스에서 추출하므로 '국문 본문'은 본문으로 처리(구형식). 키워드 줄만 예외.
            drop_orphan_caps()
            mk = re.match(r'\s*(Keywords?|주제어|키워드)\s*[:：]\s*(.*)', txt, re.I)
            if mk:
                if mk.group(1).lower().startswith('key'):
                    res['keywords_en'] = res['keywords_en'] or mk.group(2).strip()
                else:
                    res['keywords_ko'] = res['keywords_ko'] or mk.group(2).strip()
            else:
                pending.append(txt)
        elif sty == '3. 표':          # 표 캡션 — 다음 표에 연결
            pending_cap[0] = txt
        elif sty == '4. 그림' or (re.match(r'^(Figure|Fig\.?|그림)\s*\d', txt) and len(txt) < 80):
            flush()                    # 그림 캡션 → 직전 이미지에 부착, 없으면 캡션만 블록
            if res['blocks'] and res['blocks'][-1].get('type') == 'figure' and not res['blocks'][-1].get('caption'):
                res['blocks'][-1]['caption'] = txt
            else:
                res['blocks'].append({'type': 'figure', 'caption': txt, 'image': ''})
        elif sty in ('1. 대제목_국문',):
            drop_orphan_caps(); flush(); res['blocks'].append({'type': 'h1', 'text': txt})
        elif sty in ('1.1. 중제목_국문', '중제목'):
            drop_orphan_caps(); flush(); res['blocks'].append({'type': 'h2', 'text': txt})
        elif sty in ('1.1.1. 소제목_국문', '0.0. sub heading', '한1.1.1한'):
            drop_orphan_caps(); flush(); res['blocks'].append({'type': 'h3', 'text': txt})
        elif sty in ('1.1.1.1. 세부제목_국문', '0.0.0.sub heading'):
            drop_orphan_caps(); flush(); res['blocks'].append({'type': 'h4', 'text': txt})
        else:  # 2. 본문 및 기타 → 본문 (번호형 짧은 줄은 제목으로 보강 — 스타일 오류 대비)
            drop_orphan_caps()
            m = re.match(r'^(\d+(?:\.\d+){0,3})\.\s+[A-Z가-힣]', txt)
            if (m and len(txt) <= 60 and '\n' not in txt
                    and not re.search(r'[.。다음임함됨\)]\s*$', txt)):
                depth = m.group(1).count('.') + 1
                flush()
                res['blocks'].append({'type': {1: 'h1', 2: 'h2', 3: 'h3', 4: 'h4'}.get(depth, 'h4'),
                                      'text': txt})
            else:
                pending.append(txt)
    drop_orphan_caps()
    flush()
    res['title_en'] = '\n'.join(title_en).strip()
    res['title_ko'] = '\n'.join(title_ko).strip()
    res['references'] = '\n'.join(refs).strip()
    # 초록은 _extract_abstract_box(초록 박스)에서 채움 — 여기서 덮어쓰지 않음.
    if abstract and not res.get('abstract_en'):
        res['abstract_en'] = '\n'.join(a for a in abstract if not _has_hangul(a)).strip()


def _parse_generic(items, res):
    """일반 워드(Heading/번호 제목) → 최선노력 분해 (표 포함)."""
    refs, abstract, pending = [], [], []
    section = None
    title_done = False
    tbl_cap = ['']  # "Table N." 패턴 캡션 대기
    started = [False]

    def flush():
        if pending:
            res['blocks'].append({'type': 'body', 'text': '\n'.join(pending)})
            pending.clear()

    idx = 0
    for it in items:
        if it[0] == 'tbl':
            tbl = it[1]
            if _looks_like_abstract_box(tbl):
                _extract_abstract_box(tbl, res)
            elif started[0] and _table_looks_like_data(_table_tsv(tbl)):
                flush()
                res['blocks'].append({'type': 'table', 'caption': tbl_cap[0], 'data': _table_tsv(tbl)})
            tbl_cap[0] = ''
            continue
        if it[0] == 'img':                              # 이미지(그림) → 그림 블록
            flush()
            res['blocks'].append({'type': 'figure', 'caption': '', 'image': 'data:image/png;base64,' + it[1]})
            continue
        sty, txt = it[1], it[2]
        idx += 1
        if not txt:
            continue
        started[0] = True
        low = sty.lower()
        if not title_done and (low.startswith('title') or idx <= 2):
            if _has_hangul(txt) and not res['title_ko']:
                res['title_ko'] = txt
            elif not res['title_en']:
                res['title_en'] = txt
            title_done = bool(res['title_en'] or res['title_ko'])
            continue
        if re.match(r'^\s*(abstract|초록|국문\s*초록|요약)\s*$', txt, re.I):
            section = 'abstract'; flush(); continue
        if re.match(r'^\s*(references|참고\s*문헌)\s*$', txt, re.I):
            section = 'refs'; flush(); continue
        mk = re.match(r'\s*(keywords?|주제어)\s*[:：]\s*(.*)', txt, re.I)
        if mk:
            if mk.group(1).lower().startswith('key'):
                res['keywords_en'] = mk.group(2).strip()
            else:
                res['keywords_ko'] = mk.group(2).strip()
            continue
        if re.match(r'^\s*(Table|표)\s*\d', txt):   # 표 캡션 대기
            tbl_cap[0] = txt; continue
        is_heading = low.startswith('heading') or bool(re.match(r'^\d+(\.\d+)*\.?\s+\S', txt))
        if section == 'refs':
            refs.append(txt); continue
        if section == 'abstract' and not is_heading:
            abstract.append(txt); continue
        if is_heading:
            section = 'body'; flush()
            level = 'h2' if (low == 'heading 2' or re.match(r'^\d+\.\d+', txt)) else 'h1'
            res['blocks'].append({'type': level, 'text': txt})
            continue
        section = 'body'; pending.append(txt)
    flush()
    res['references'] = '\n'.join(refs).strip()
    if abstract and not res.get('abstract_en'):
        res['abstract_en'] = '\n'.join(a for a in abstract if not _has_hangul(a)).strip()


def _block_nonempty(b):
    if b['type'] == 'table':
        return bool(b.get('caption', '').strip() or b.get('data', '').strip())
    if b['type'] == 'figure':
        return bool(b.get('caption', '').strip() or b.get('image'))
    return bool(b.get('text', '').strip())


def _extract_footnotes(file_bytes):
    """불러온 docx의 footnotes.xml에서 저자 소속 각주 텍스트 추출(저작권 제외)."""
    import zipfile
    try:
        z = zipfile.ZipFile(io.BytesIO(file_bytes))
        if 'word/footnotes.xml' not in z.namelist():
            return []
        x = z.read('word/footnotes.xml').decode('utf-8')
        out = []
        import html as _html
        for fid, body in re.findall(r'<w:footnote[^>]*w:id="(\d+)"[^>]*>(.*?)</w:footnote>', x, re.S):
            txt = _html.unescape(''.join(re.findall(r'<w:t[^>]*>([^<]*)</w:t>', body))).strip()
            # 소속 뒤에 붙은 저작권 문구 제거(중복 방지)
            txt = re.split(r'[ⓒ©]\s*Copyright|Copyright\s*:\s*The Author', txt)[0].strip()
            if txt and not re.match(r'^[ⓒ©]?\s*Copyright', txt, re.I):
                out.append(txt)
        return out
    except Exception:
        return []


def parse_docx(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    items = list(_iter_body_items(doc))
    res = {k: '' for k in ['title_en', 'title_ko', 'authors', 'received', 'revised',
                           'accepted', 'abstract_en', 'keywords_en', 'abstract_ko',
                           'keywords_ko', 'jel', 'references']}
    res['blocks'] = []
    # 구형 포맷: 제목·저자·초록이 상단 표 안에 있음 → 먼저 추출하고 그 표는 본문 파서에서 제외
    kept = []
    header_done = False
    for it in items:
        if it[0] == 'tbl' and not header_done and _is_old_header_table(it[1]):
            _extract_header_table(it[1], res)
            header_done = True
            continue
        kept.append(it)
    items = kept
    styles = {it[1] for it in items if it[0] == 'p'}
    # KJFM 양식 신호: 표준 스타일 / 저자 스타일 / 초록 박스가 하나라도 있으면 위치기반 _parse_kjfm
    has_kjfm = bool({'Title_English', '2. 본문', '1. 대제목_국문', '국문제목'} & styles)
    has_author = bool(set(_AUTHOR_STYLES) & styles)
    has_absbox = any(it[0] == 'tbl' and _looks_like_abstract_box(it[1]) for it in items)
    if has_kjfm or has_author or has_absbox:
        _parse_kjfm(items, res)
    else:
        _parse_generic(items, res)
    res['blocks'] = [b for b in res['blocks'] if _block_nonempty(b)]
    # 각주 분류: 소속(Affiliation/교수/대학/Email)은 affiliations, 그 외(지원·감사)는 제목 설명각주
    _foots = _extract_footnotes(file_bytes)
    _affs, _tnote = [], ''
    for f in _foots:
        if re.search(r'Affiliation|Professor|University|Email|소속|교수|대학|college|Dept', f, re.I):
            _affs.append(f)
        elif not _tnote:
            _tnote = f
        else:
            _affs.append(f)
    res['affiliations'] = '\n'.join(_affs)
    res['title_note'] = _tnote
    # DOI 추출(상단 표) → 칸에 채워 편집 가능하게. 'KJFM.' 뒤 부분만.
    if not res.get('doi'):
        for tbl in doc.tables[:2]:
            m = re.search(r'KJFM\.([0-9][0-9.]*)', _table_cell_text(tbl))
            if m:
                res['doi'] = m.group(1).rstrip('.'); break
    # DOI 규칙 {연}.{월}.{권}.{호}.{시작페이지} → 권(호)·시작페이지 칸 채움
    _dn = re.findall(r'\d+', res.get('doi', '') or '')
    if len(_dn) >= 5:
        res.setdefault('vol_issue', ''); res.setdefault('start_page', '')
        if not res['vol_issue']:
            res['vol_issue'] = '%s(%s)' % (_dn[2], _dn[3])
        if not res['start_page']:
            res['start_page'] = _dn[4]
    # 언어 감지: 본문이 한글 위주면 국문 논문, 아니면 영문 논문
    body_text = ' '.join(b.get('text', '') for b in res['blocks'][:6])
    ko = len(re.findall(r'[가-힣]', body_text)); en = len(re.findall(r'[A-Za-z]', body_text))
    res['lang'] = 'ko' if ko >= en else 'en'
    return res


# ---------------------------------------------------------------- HTTP 핸들러
